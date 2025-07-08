from docx import Document
import sys
import speech_recognition as sr
import pyttsx3
import datetime
import webbrowser
import os
from together import Together
import time
import random
import re
import subprocess
import difflib
from dotenv import load_dotenv

# ------------------- SETUP -------------------
load_dotenv()
together_api_key = os.getenv("TOGETHER_API_KEY")

if not together_api_key:
    print("Error: TOGETHER_API_KEY is not set.")
    sys.exit("API key missing. Exiting...")

client = Together(api_key=together_api_key)

# ------------------- TEXT TO SPEECH -------------------
engine = pyttsx3.init(driverName='sapi5')

def say(text):
    engine.say(text)
    engine.runAndWait()

def list_available_voices():
    voices = engine.getProperty('voices')
    for voice in voices:
        gender = getattr(voice, "gender", "Unknown")
        print(f"Name: {voice.name}, ID: {voice.id}, Gender: {gender}")

def set_voice_by_name(name: str):
    name = name.lower().strip()
    name_aliases = {
        "jeera": "zira", "jeeera": "zira", "zira": "zira", "david": "david"
    }
    name = name_aliases.get(name, name)

    voices = engine.getProperty('voices')
    for voice in voices:
        if name in voice.name.lower():
            engine.setProperty('voice', voice.id)
            say(f"Hi, I am {voice.name}, at your service.")
            return
    say("Sorry, I couldn't find that voice.")

# ------------------- AI FUNCTION -------------------
def ask_me(prompt):
    try:
        response = client.chat.completions.create(
            model="deepseek-ai/DeepSeek-V3",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1024,
            temperature=0.7
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"[Together API Error]: {e}")
        return "Sorry, I couldn't get a response."

# ------------------- SPEECH RECOGNITION -------------------
def takeCommand():
    r = sr.Recognizer()
    with sr.Microphone() as source:
        r.adjust_for_ambient_noise(source)
        print("Listening...")
        try:
            audio = r.listen(source, timeout=5)
            query = r.recognize_google(audio, language="en-in")
            print(f"User said: {query}")
            return query.lower()
        except sr.UnknownValueError:
            say("Sorry, I didn't catch that.")
        except sr.RequestError as e:
            say("Speech service is not working.")
            print(f"Request error: {e}")
        except Exception as e:
            print(f"Recognition error: {e}")
    return ""

# ------------------- DOCUMENT CREATION -------------------
def create_word_doc_from_topic(topic):
    say(f"Creating article for {topic}")
    content = ask_me(f"Write a detailed article about {topic}")
    try:
        doc = Document()
        doc.add_heading(f"Topic: {topic}", 0)
        doc.add_paragraph(content)

        safe_topic = re.sub(r'[\\/*?:"<>|]', "_", topic)
        filename = f"{safe_topic.replace(' ', '_')}.docx"
        doc.save(filename)

        say(f"Document saved as {filename}")
    except Exception as e:
        print(f"Error creating document: {e}")
        say("Failed to create the document.")

# ------------------- GAME: NUMBER GUESSING -------------------
def number_guessing_game():
    say("Let's play a number guessing game.")
    number = random.randint(1, 20)
    for _ in range(5):
        say("Take a guess.")
        guess = takeCommand()
        if guess.isdigit():
            guess = int(guess)
            if guess < number:
                say("Too low.")
            elif guess > number:
                say("Too high.")
            else:
                say("You got it!")
                return
        else:
            say("Please say a number.")
    say(f"The number was {number}. Better luck next time!")

# ------------------- GAME: TRUTH OR DARE -------------------
truths = ["What is your biggest fear?", "What's a secret you've never told?", "Have you ever lied to a friend?", "What's your most embarrassing moment?", "Who was your first crush?"]
dares = ["Do 10 pushups.", "Sing your favorite song chorus.", "Dance like a chicken.", "Text someone 'I like you'.", "Do an impression of your favorite actor."]

def get_players():
    players = []
    print("Enter player names (type 'done' to finish):")
    while True:
        name = input("Player name: ").strip()
        if name.lower() == 'done':
            break
        if name:
            players.append(name)
    return players

def play_game(players):
    print("\n🎉 Starting Truth or Dare! 🎉\n")
    while True:
        player = random.choice(players)
        print(f"\n👉 {player}'s turn!")
        choice = input("Type 'truth', 'dare', or 'quit': ").strip().lower()
        if choice == 'truth':
            print("🧠 Truth:", random.choice(truths))
        elif choice == 'dare':
            print("🎯 Dare:", random.choice(dares))
        elif choice == 'quit':
            print("Thanks for playing!")
            break
        else:
            print("Invalid input.")
        time.sleep(1.5)

# ------------------- MUSIC PLAYBACK -------------------
def list_songs_in_downloads():
    path = os.path.join(os.path.expanduser("~"), "Downloads")
    audio_ext = ('.mp3', '.aac', '.ogg', '.m4a', '.wav')
    return [f for f in os.listdir(path) if f.endswith(audio_ext)], path

def play_song(path):
    try:
        if sys.platform == 'win32':
            os.startfile(path)
        elif sys.platform == 'darwin':
            subprocess.call(['open', path])
        else:
            subprocess.call(['xdg-open', path])
    except Exception as e:
        say("Could not play the song.")
        print("Error:", e)

def choose_and_play_song():
    songs, path = list_songs_in_downloads()
    if not songs:
        say("No songs found.")
        return

    say("Which song do you want to play?")
    query = takeCommand()
    match = difflib.get_close_matches(query, songs, n=1, cutoff=0.7)

    if match:
        selected = os.path.join(path, match[0])
        say(f"Playing {match[0]}")
        play_song(selected)
    else:
        say("Couldn't find that. Here's the list.")
        for idx, song in enumerate(songs):
            print(f"{idx + 1}. {song}")
            say(f"{idx + 1}")
        say("Say the number.")
        choice = takeCommand()
        if choice.isdigit():
            i = int(choice)
            if 1 <= i <= len(songs):
                play_song(os.path.join(path, songs[i - 1]))
            else:
                say("Invalid number.")
        else:
            say("Invalid input.")

# ------------------- MAIN FUNCTION -------------------
def run_Xebec():
    say("Hello, I am Xebec. I'm awake.")
    while True:
        query = takeCommand()
        if not query:
            continue

        if any(cmd in query for cmd in ["stop", "exit", "quit", "shut down"]):
            say("Do you want me to shut down?")
            confirm = takeCommand()
            if any(w in confirm for w in ["yes", "sure", "okay"]):
                say("Shutting down...")
                return "stop"
            else:
                say("Shutdown canceled.")

        elif "play song" in query or "play music" in query:
            choose_and_play_song()

        elif "switch to" in query and "voice" in query:
            name = query.replace("switch to", "").replace("voice", "").strip()
            if name:
                set_voice_by_name(name)
            else:
                say("Please say the name.")

        elif "create document" in query or "write article" in query:
            say("What topic?")
            topic = takeCommand()
            if topic:
                create_word_doc_from_topic(topic)
            else:
                say("I didn't hear the topic.")

        elif "the time" in query:
            time_now = datetime.datetime.now().strftime('%I:%M %p')
            say(f"The time is {time_now}")

        elif "open" in query:
            sites = {
                "youtube": "https://www.youtube.com",
                "instagram": "https://www.instagram.com",
                "facebook": "https://www.facebook.com",
                "google": "https://www.google.com",
                "chat gpt": "https://www.chatgpt.com",
                "spotify": "https://www.spotify.com",
                "github": "https://www.github.com",
                "deepseek": "https://www.deepseek.com",
                "canva": "https://www.canva.com",
                "whatsapp": "https://www.whatsapp.com",
                "discord": "https://www.discord.com",
            }
            opened = False
            for site in sites:
                if f"open {site}" in query:
                    say(f"Opening {site}")
                    webbrowser.open(sites[site])
                    opened = True
                    break
            if not opened:
                say("Site not found in my list.")

        elif "truth or dare" in query:
            say("Starting the game.")
            players = get_players()
            if players:
                play_game(players)
            else:
                say("No players. Ending.")

        elif "play game" in query:
            say("Which game? Say 'number guessing'.")
            choice = takeCommand()
            if "number" in choice:
                number_guessing_game()
            else:
                say("Only number guessing is available.")

        elif any(x in query for x in ["who is", "what is", "define", "tell me", "ask me", "where is"]):
            response = ask_me(query)
            print("Xebec:", response)
            say(response)

# ------------------- START -------------------
if __name__ == '__main__':
    while True:
        result = run_Xebec()
        if result == "stop":
            say("Goodbye.")
            break
